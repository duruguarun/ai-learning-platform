import requests
from docx import Document

url = "http://127.0.0.1:8000/analyse"

files = [
    ("files", open(r"C:\Users\venky\Downloads\Data Structure-1.pdf", "rb"))
]

response = requests.post(url, files=files)
data = response.json()

# Create Word document
doc = Document()
doc.add_heading("Exam Paper Analysis Report", level=1)

# Frequent Questions
doc.add_heading("Frequent Questions", level=2)
for q in data.get("frequent_questions", []):
    doc.add_paragraph(q, style="List Bullet")

# Ranked Topics
doc.add_heading("Ranked Topics", level=2)
for topic in data.get("ranked_topics", []):
    doc.add_paragraph(str(topic), style="List Bullet")

# Paper Pattern
doc.add_heading("Paper Pattern", level=2)
doc.add_paragraph(str(data.get("paper_pattern", "")))

# Important Topics Summary
doc.add_heading("Important Topics Summary", level=2)
doc.add_paragraph(str(data.get("important_topics_summary", "")))

# Save file
doc.save("analysis_report.docx")

print("✅ Word file created: analysis_report.docx")